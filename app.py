import re
import os
from datetime import datetime
from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS # <--- TAMBAHKAN INI
from docx import Document
from docx.shared import Pt

app = Flask(__name__, template_folder='.')
CORS(app) # <--- TAMBAHKAN INI (Mengizinkan akses dari GitHub Pages)

app = Flask(__name__, template_folder='.')

def bersihkan_dan_urai_kresna(text):
    """
    ENGINE PARSER V5.1: Memisahkan komponen teks secara akurat.
    """
    matches = list(re.finditer(r"\bBidang\s*:", text))
    if matches:
        text_clean = text[matches[-1].start():]
    else:
        text_clean = text

    text_clean = re.sub(r'\[[^\]]+\]\s*KRESNA AI:\s*', '', text_clean)
    text_clean = re.sub(r'KRESNA AI:\s*', '', text_clean)

    data_terurai = {}

    def ambil_pola(pattern, src, flags=0):
        match = re.search(pattern, src, flags)
        return match.group(1).strip() if match else ""

    # GENERATE WAKTU REALTIME
    now = datetime.now()
    hari_list = ["Senin", "Selasa", "Rabu", "Kamis", "Jumat", "Sabtu", "Minggu"]
    bulan_list = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
    
    hari_ini = hari_list[now.weekday()]
    tgl = now.day
    bln = bulan_list[now.month]
    thn = now.year
    jam_menit = now.strftime("%H.%M")

    data_terurai['{{WAKTU}}'] = f"{hari_ini}, {tgl} {bln} {thn} Pukul {jam_menit} WIB"
    data_terurai['{{TANGGAL_TTD}}'] = f"{tgl} {bln} {thn}"

    data_terurai['{{BIDANG}}'] = ambil_pola(r"Bidang\s*:\s*(.*)", text_clean)
    data_terurai['{{PERIHAL}}'] = ambil_pola(r"Perihal\s*:\s*(.*)", text_clean)
    data_terurai['{{SUMBER}}'] = ambil_pola(r"A\.\s*Sumber Informasi\s*:\s*(.*)", text_clean)
    data_terurai['{{HUBUNGAN}}'] = ambil_pola(r"B\.\s*Hubungan dengan Sumber\s*:\s*(.*)", text_clean)
    data_terurai['{{CARA}}'] = ambil_pola(r"C\.\s*Cara mendapatkan Informasi\s*:\s*(.*)", text_clean)
    data_terurai['{{NILAI}}'] = ambil_pola(r"E\.\s*Nilai Informasi\s*:\s*(.*)", text_clean)

    fakta_block = ambil_pola(r"II\.\s*FAKTA-FAKTA(.*?)III\.\s*PENDAPAT", text_clean, re.DOTALL)
    if fakta_block:
        data_terurai['{{FAKTA_A}}'] = ambil_pola(r"A\.\s*(.*?)(?=B\.\s*|$)", fakta_block, re.DOTALL)
        data_terurai['{{FAKTA_B}}'] = ambil_pola(r"B\.\s*(.*?)(?=C\.\s*|$)", fakta_block, re.DOTALL)
        data_terurai['{{FAKTA_C}}'] = ambil_pola(r"C\.\s*(.*)", fakta_block, re.DOTALL)
    else:
        data_terurai['{{FAKTA_A}}'] = data_terurai['{{FAKTA_B}}'] = data_terurai['{{FAKTA_C}}'] = ""

    pendapat_block = ambil_pola(r"III\.\s*PENDAPAT PELAPOR(.*)", text_clean, re.DOTALL)
    if pendapat_block:
        abjad = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
        for idx, char in enumerate(abjad):
            next_char = abjad[idx+1] if idx+1 < len(abjad) else None
            pattern = rf"{char}\.\s*(.*?)(?={next_char}\.\s*|$)" if next_char else rf"{char}\.\s*(.*)"
            data_terurai[f'{{{{PENDAPAT_{char}}}}}'] = ambil_pola(pattern, pendapat_block, re.DOTALL)
    else:
        for char in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']:
            data_terurai[f'{{{{PENDAPAT_{char}}}}}'] = ""

    return data_terurai

def suntik_dokumen_docx(p, data_komponen):
    p_text = p.text
    ada_perubahan = False
    
    if "Waktu mendapatkan Informasi" in p_text and ":" in p_text:
        prefix = p_text.split(":")[0]
        p_text = f"{prefix}: {data_komponen['{{WAKTU}}']}"
        ada_perubahan = True
    elif "Bandar Lampung" in p_text:
        p_text = f"Bandar Lampung, {data_komponen['{{TANGGAL_TTD}}']}"
        ada_perubahan = True
    else:
        for key, val in data_komponen.items():
            if key in p_text:
                p_text = p_text.replace(key, val)
                ada_perubahan = True
                
    if ada_perubahan:
        is_bold = p.runs[0].bold if p.runs else False
        p.text = "" 
        run = p.add_run(p_text)
        run.font.name = 'Tahoma'
        run.bold = is_bold
        run.font.size = Pt(11)

# ================= ROUTE API =================

@app.route('/')
def index():
    return render_template('https://kresnaai.github.io/LI-docx/dashboard.html')

@app.route('/generate', methods=['POST'])
@app.route('/api/generate', methods=['POST'])
def generate_doc():
    data = request.json
    narasi = data.get('text', '')
    storage_path = data.get('storage_path', '').strip()

    if not narasi:
        return jsonify({"status": "error", "message": "Input kosong!"}), 400
        
    komponen_data = bersihkan_dan_urai_kresna(narasi)
    master_docx = "format_original.docx"
    output_lokal = "output_LI.docx"
    
    try:
        doc = Document(master_docx)
        
        for p in doc.paragraphs:
            suntik_dokumen_docx(p, komponen_data)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        suntik_dokumen_docx(p, komponen_data)
                        
        # Simpan backup lokal utama
        doc.save(output_lokal)
        
        # Pengecekan Sistem Operasi (Windows vs WSL/Linux)
        if storage_path:
            if os.name == 'nt':  # JIKA DIJALANKAN DI WINDOWS (PowerShell)
                storage_path = os.path.normpath(storage_path)
            else:  # JIKA DIJALANKAN DI LINUX / WSL
                if re.match(r'^[a-zA-Z]:', storage_path):
                    drive = storage_path[0].lower()
                    rem_path = storage_path[2:].replace('\\', '/')
                    storage_path = f"/mnt/{drive}{rem_path}"
                else:
                    storage_path = storage_path.replace('\\', '/')
            
            # Buat folder jika belum ada dan simpan file di sana
            os.makedirs(storage_path, exist_ok=True)
            custom_output_target = os.path.join(storage_path, "Laporan_Informasi_Kresna.docx")
            doc.save(custom_output_target)
            
        return jsonify({"status": "success", "docx_url": "/download/docx"})
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/download/docx')
def download_docx():
    return send_file("output_LI.docx", as_attachment=True, download_name="Laporan_Informasi_Kresna.docx")

if __name__ == '__main__':
    app.run(debug=True, port=5000)
